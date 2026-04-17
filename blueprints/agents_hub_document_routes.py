"""Artifact-generation route (hub_generate_document) for agents_hub_bp —
pptx/docx/xlsx/csv/markdown/pdf/infographic generation. Split out of
agents_hub_bp.py in Phase 3 Slice 4 (kept as its own file given its size).
"""

import sys, uuid, json, logging, threading, re, calendar
from datetime import datetime, timedelta, date
from flask import (render_template, jsonify, request,
                   Response, stream_with_context, redirect, url_for, abort)
import auth
import token_limits
from agents_hub_bp import agents_hub_bp
from agents_hub_bp import (
    _ss_exec, _fix_row, _fix_rows, _is_dev_or_admin, _is_admin,
    _can_access_agent, _agent_dict,
)
from llm_providers.factory import get_provider
from llm_providers.models import resolve_default_model


@agents_hub_bp.route('/api/agenthub/generate-document', methods=['POST'])
@auth.login_required
def hub_generate_document():
    """Generate a downloadable document from free-form conversation text.

    No SQL rows required — works with any agent (web search, tools, etc.)
    purely from the chat transcript. Supports ``type`` values: "presentation"
    (pptx slide JSON via InfographicGenerator-style layout), "infographic",
    "docx", "csv", "xlsx", "markdown", "txt", "pdf". Each branch prompts the
    LLM to extract/structure data from the conversation, then builds the
    actual file (or, for presentation/infographic, a JSON layout the
    frontend renders).

    POST body:
    {
        "type":        "presentation" | "infographic" | "docx" | "csv" | "xlsx" | "markdown" | "txt" | "pdf",
        "messages":    ["User: ...", "Assistant: ..."],
        "instruction": "make me a ppt with 4 slides...",
        "agent_name":  "Agent Name"
    }

    Returns:
        JSON with ``status``, and depending on type either ``slides``/
        ``infographic`` (rendered client-side) or ``filename``, ``mime_type``,
        ``content_b64``/``content_text``, ``preview_html``, ``size_bytes``.
    """
    import os, re as _re

    data        = request.get_json(force=True) or {}
    doc_type    = data.get('type', 'presentation')
    messages    = data.get('messages', [])
    instruction = data.get('instruction', '')
    agent_name  = data.get('agent_name', '')

    if not messages:
        return jsonify({"status": "error", "message": "messages is required"}), 400

    user = auth.current_user()

    # ── Token gate ────────────────────────────────────────────────────────────
    try:
        import token_limits as _tl
        allowed, msg, _, _ = _tl.check_limit(user)
        if not allowed:
            return jsonify({"status": "error", "message": msg}), 429
    except ImportError:
        pass

    conv_text = '\n'.join(messages)

    # Extract a short topic label from the conversation for labelling slides
    def _topic_from_messages(msgs, instr):
        combined = ' '.join(msgs[:4] + ([instr] if instr else []))
        nouns = _re.findall(r'\b([A-Z][a-zA-Z]{2,})\b', combined)
        skip  = {'User','Assistant','Sure','Hello','Here','Please','Let',
                 'The','This','That','Yes','No','Thanks','Great','Okay'}
        for n in nouns:
            if n not in skip:
                return n
        return agent_name or 'Conversation'

    topic = _topic_from_messages(messages, instruction)

    # ── Resolve LLM provider ─────────────────────────────────────────────────
    try:
        _llm_provider = get_provider()
    except Exception as e:
        return jsonify({"status": "error", "message": f"LLM client unavailable: {e}"}), 500

    model = resolve_default_model("OPENAI_MODEL", tier="quality")

    # ─────────────────────────────────────────────────────────────────────────
    # PRESENTATION
    # ─────────────────────────────────────────────────────────────────────────
    if doc_type == 'presentation':
        # How many slides did the user ask for?
        m = _re.search(r'\b(\d+)\s+slides?\b', instruction, _re.IGNORECASE)
        n_slides = int(m.group(1)) if m else 4
        n_slides = max(2, min(n_slides, 8))

        themes = ['delivery', 'process', 'competitive']
        slide_themes = [themes[i % 3] for i in range(n_slides)]

        prompt = f"""You are a world-class presentation designer. Your slides must be visually rich, data-driven, and tell a clear story.

USER REQUEST: {instruction or f'Create a compelling {n_slides}-slide presentation from this conversation.'}

CONVERSATION:
{conv_text[:7000]}

──────────────────────────────────────────────────────
OUTPUT RULES (read carefully — violations will be rejected):
1. Return ONLY a raw JSON array. No markdown fences, no explanations, no text before/after.
2. The array must contain EXACTLY {n_slides} slide objects.
3. Every slide must have ALL these fields.

SLIDE SCHEMA:
{{
  "theme":          "delivery" | "process" | "competitive",
  "title":          "Bold, specific slide title (e.g. 'Bitcoin Hits $65K — Bull Run Confirmed')",
  "subtitle":       "One crisp sentence that sets up the slide",
  "headline_stat":  "The single most impactful number on this slide (e.g. '$65,420', '+127%', '3.2T')",
  "headline_label": "Plain-English label for that number (e.g. 'BTC Price Today', 'YTD Return', 'Market Cap')",
  "bottom_message": "One actionable takeaway or insight the audience should remember",
  "columns": [
    {{
      "col_title":    "Column heading (short, specific)",
      "col_subtitle": "Column sub-heading providing context",
      "bullets": ["Specific fact or number", "Another concrete insight", "Third supporting point"]
    }},
    {{
      "col_title":    "...", "col_subtitle": "...",
      "bullets": ["...", "...", "..."]
    }},
    {{
      "col_title":    "...", "col_subtitle": "...",
      "bullets": ["...", "...", "..."]
    }}
  ]
}}

QUALITY RULES:
• headline_stat MUST always be present — pick the single most eye-catching number for every slide.
• Bullets must be SPECIFIC: use exact figures, dates, percentages from the conversation.
  BAD:  "Price went up significantly"
  GOOD: "Q1 2022: $47,500 → Q2 2022: $20,300 (-57%)"
• Each column should cover a distinct angle (e.g. Slide about trends → Price / Volume / Sentiment).
• Theme assignment: {', '.join(f'slide {i+1}: {t}' for i, t in enumerate(slide_themes))}
• Slide 1: headline information (current state, key figure). Remaining slides: depth/breakdown/trends.
• Max 120 characters per bullet. Exactly 3 bullets per column."""

        try:
            result = _llm_provider.chat(
                system      = "",
                messages    = [{"role": "user", "content": prompt}],
                model       = model,
                temperature = 0.25,
                max_tokens  = 4000,
            )
            raw = (result.content or "").strip()

            # Extract JSON array robustly
            arr_match = _re.search(r'\[[\s\S]*\]', raw)
            if arr_match:
                slides = json.loads(arr_match.group(0))
            else:
                parsed = json.loads(raw)
                slides = parsed if isinstance(parsed, list) else list(parsed.values())[0]

            if not isinstance(slides, list) or not slides:
                raise ValueError("No slides in response")

            # Normalise each slide
            for i, s in enumerate(slides):
                if s.get('theme') not in themes:
                    s['theme'] = themes[i % 3]
                cols = s.get('columns')
                if not isinstance(cols, list):
                    cols = []
                # Ensure exactly 3 columns
                while len(cols) < 3:
                    cols.append({"col_title": "Details", "col_subtitle": "", "bullets": ["—", "—", "—"]})
                s['columns'] = cols[:3]
                for col in s['columns']:
                    b = col.get('bullets', [])
                    if not isinstance(b, list):
                        b = []
                    while len(b) < 3:
                        b.append("—")
                    col['bullets'] = [str(x)[:130] for x in b[:3]]
                    s['columns'] = s.get('columns', cols[:3])

            # Token recording
            _in_tok       = result.usage.prompt_tokens
            _out_tok      = result.usage.completion_tokens
            tokens_used   = result.usage.total_tokens
            try:
                import token_limits as _tl2
                if user and user.get("id"):
                    _tl2.record_usage(
                        user_id       = user["id"],
                        tokens        = tokens_used,
                        call_type     = "hub_presentation",
                        agent_name    = topic,
                        question      = (instruction or "Generate presentation")[:200],
                        input_tokens  = _in_tok,
                        output_tokens = _out_tok,
                        model         = model,
                    )
            except Exception:
                pass

            return jsonify({
                "status":      "success",
                "slides":      slides,
                "client":      {"name": topic},
                "client_id":   "",
                "pptx_base64": None,
            })

        except Exception as exc:
            import traceback; traceback.print_exc()
            return jsonify({"status": "error", "message": str(exc)}), 500

    # ─────────────────────────────────────────────────────────────────────────
    # INFOGRAPHIC  (also handles "report" from the frontend)
    # ─────────────────────────────────────────────────────────────────────────
    elif doc_type == 'infographic':
        try:
            from infographicgenerator import InfographicGenerator

            # Use GPT to extract the best, most specific summary points
            extract_prompt = f"""You are extracting key insights from a conversation for an infographic.

CONVERSATION:
{conv_text[:5000]}

USER REQUEST: {instruction or 'Create a visual infographic summary'}

Extract 10-15 highly specific, data-rich insight bullets. Each bullet should:
- Be concise (under 90 chars)
- Contain actual numbers, dates, or percentages where available
- Read as a standalone fact (no "we discussed" or conversational language)
- Cover the most important points across the whole conversation

Return ONLY a JSON array of strings. Example:
["Bitcoin current price: $65,420 (+2.3% in 24h)", "Q1 2022 low: $17,600 (-62% from peak)", ...]"""

            try:
                ext_result = _llm_provider.chat(
                    system      = "",
                    messages    = [{"role": "user", "content": extract_prompt}],
                    model       = model,
                    temperature = 0.2,
                    max_tokens  = 1000,
                )
                raw_pts = (ext_result.content or "").strip()
                arr_m   = _re.search(r'\[[\s\S]*\]', raw_pts)
                summary_points = json.loads(arr_m.group(0)) if arr_m else json.loads(raw_pts)
                if not isinstance(summary_points, list):
                    raise ValueError("not a list")
                summary_points = [str(p) for p in summary_points if p]
            except Exception:
                # Fallback: extract sentences from assistant messages
                summary_points = []
                for msg in messages:
                    if msg.startswith('Assistant:'):
                        body = msg[len('Assistant:'):].strip()
                        for s in _re.split(r'(?<=[.!?])\s+', body):
                            if len(s.strip()) > 20:
                                summary_points.append(s.strip())
                if not summary_points:
                    summary_points = [conv_text[:2000]]

            gen    = InfographicGenerator()
            layout = gen.generate_infographic_layout(
                summary_points = summary_points[:18],
                rows           = [],
                user           = user,
                agent_name     = agent_name or topic,
            )
            if not layout:
                return jsonify({"status": "error", "message": "Infographic generation failed"}), 400
            return jsonify({"status": "success", "infographic": layout})

        except Exception as exc:
            import traceback; traceback.print_exc()
            return jsonify({"status": "error", "message": str(exc)}), 500

    # ─────────────────────────────────────────────────────────────────────────
    # WORD DOCUMENT (.docx)
    # ─────────────────────────────────────────────────────────────────────────
    elif doc_type == 'docx':
        try:
            from docx import Document as _DocxDoc
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io as _io, base64 as _b64

            struct_prompt = f"""You are a professional document writer. Create a Word document from the conversation below.

USER REQUEST: {instruction or 'Write a comprehensive Word document summarising the conversation.'}

CONVERSATION:
{conv_text[:7000]}

Return ONLY a JSON object with this exact structure (no markdown fences):
{{
  "title": "Document Title",
  "subtitle": "Optional brief subtitle or date",
  "sections": [
    {{"type": "heading1", "text": "Section Title"}},
    {{"type": "paragraph", "text": "Paragraph content. Use real data and specifics."}},
    {{"type": "heading2", "text": "Sub-section"}},
    {{"type": "bullet_list", "items": ["Specific point 1", "Point 2", "Point 3"]}},
    {{"type": "table", "headers": ["Col A", "Col B", "Col C"], "rows": [["r1a","r1b","r1c"]]}}
  ]
}}

Rules:
- Use REAL data, numbers and facts from the conversation
- Structure logically (summary → details → conclusions)
- Include a table when tabular data exists in the conversation
- Minimum 6 sections, maximum 20"""

            result = _llm_provider.chat(
                system="", model=model, temperature=0.2, max_tokens=3000,
                messages=[{"role": "user", "content": struct_prompt}],
            )
            raw = (result.content or "").strip()
            arr_m = _re.search(r'\{[\s\S]*\}', raw)
            struct = json.loads(arr_m.group(0) if arr_m else raw)

            # Build docx
            doc = _DocxDoc()

            # Title
            title_para = doc.add_heading(struct.get('title', topic), level=0)
            title_para.runs[0].font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

            if struct.get('subtitle'):
                sub = doc.add_paragraph(struct['subtitle'])
                sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
                sub.runs[0].font.size = Pt(11)

            doc.add_paragraph()  # spacer

            # Build preview HTML at the same time
            preview_parts = [f'<h1>{_html_esc(struct.get("title", topic))}</h1>']
            if struct.get('subtitle'):
                preview_parts.append(f'<p style="color:#64748b;font-size:.9rem">{_html_esc(struct["subtitle"])}</p>')

            for sec in struct.get('sections', []):
                t = sec.get('type', '')
                if t == 'heading1':
                    doc.add_heading(sec.get('text', ''), level=1)
                    preview_parts.append(f'<h2>{_html_esc(sec.get("text",""))}</h2>')
                elif t == 'heading2':
                    doc.add_heading(sec.get('text', ''), level=2)
                    preview_parts.append(f'<h3>{_html_esc(sec.get("text",""))}</h3>')
                elif t == 'paragraph':
                    doc.add_paragraph(sec.get('text', ''))
                    preview_parts.append(f'<p>{_html_esc(sec.get("text",""))}</p>')
                elif t == 'bullet_list':
                    items = sec.get('items', [])
                    lis = ''.join(f'<li>{_html_esc(i)}</li>' for i in items)
                    preview_parts.append(f'<ul>{lis}</ul>')
                    for item in items:
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(str(item))
                elif t == 'table':
                    hdrs = sec.get('headers', [])
                    rows_data = sec.get('rows', [])
                    if hdrs and rows_data:
                        tbl = doc.add_table(rows=1 + len(rows_data), cols=len(hdrs))
                        tbl.style = 'Table Grid'
                        for j, h in enumerate(hdrs):
                            tbl.rows[0].cells[j].text = str(h)
                        for i, row in enumerate(rows_data):
                            for j, val in enumerate(row[:len(hdrs)]):
                                tbl.rows[i+1].cells[j].text = str(val)
                        # HTML table
                        ths = ''.join(f'<th>{_html_esc(h)}</th>' for h in hdrs)
                        trs = ''.join(
                            '<tr>' + ''.join(f'<td>{_html_esc(v)}</td>' for v in row[:len(hdrs)]) + '</tr>'
                            for row in rows_data
                        )
                        preview_parts.append(f'<table><thead><tr>{ths}</tr></thead><tbody>{trs}</tbody></table>')

            buf = _io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            b64 = _b64.b64encode(buf.read()).decode()
            preview_html = '\n'.join(preview_parts)

            fname = _re.sub(r'[^\w\-]', '-', topic.lower())[:30] + '.docx'
            return jsonify({
                "status":      "success",
                "doc_type":    "docx",
                "filename":    fname,
                "mime_type":   "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "content_b64": b64,
                "content_text": None,
                "preview_html": preview_html,
                "size_bytes":  len(buf.getbuffer()) if hasattr(buf, 'getbuffer') else len(b64)*3//4,
            })

        except Exception as exc:
            import traceback; traceback.print_exc()
            return jsonify({"status": "error", "message": str(exc)}), 500

    # ─────────────────────────────────────────────────────────────────────────
    # CSV
    # ─────────────────────────────────────────────────────────────────────────
    elif doc_type == 'csv':
        try:
            import csv as _csv, io as _io

            csv_prompt = f"""Extract all tabular / structured data from the conversation and convert it to CSV.

USER REQUEST: {instruction or 'Create a CSV file from the data in this conversation.'}

CONVERSATION:
{conv_text[:6000]}

Return ONLY a JSON object (no markdown fences):
{{
  "filename_hint": "short-descriptive-name",
  "headers": ["Column1", "Column2", ...],
  "rows": [
    ["val1", "val2", ...],
    ...
  ]
}}

Rules:
- Column headers must be clear and specific
- Include ALL numerical/factual data from the conversation
- If data has dates, include a Date/Period column
- Minimum 3 columns, at least 2 data rows (create realistic rows if conversation only has a few data points)"""

            result = _llm_provider.chat(
                system="", model=model, temperature=0.15, max_tokens=2000,
                messages=[{"role": "user", "content": csv_prompt}],
            )
            raw  = (result.content or "").strip()
            arr_m = _re.search(r'\{[\s\S]*\}', raw)
            struct = json.loads(arr_m.group(0) if arr_m else raw)

            headers   = struct.get('headers', [])
            rows_data = struct.get('rows', [])
            hint      = struct.get('filename_hint', topic.lower())

            buf = _io.StringIO()
            writer = _csv.writer(buf)
            writer.writerow(headers)
            writer.writerows(rows_data)
            csv_text = buf.getvalue()

            # Build preview HTML table
            ths = ''.join(f'<th>{_html_esc(h)}</th>' for h in headers)
            trs = ''.join(
                '<tr><td class="rn">' + str(i+1) + '</td>' +
                ''.join(f'<td>{_html_esc(v)}</td>' for v in row) + '</tr>'
                for i, row in enumerate(rows_data)
            )
            preview_html = f'''
<div class="table-card">
  <div class="table-meta">
    <span><strong>{len(headers)}</strong> columns</span>
    <span><strong>{len(rows_data)}</strong> rows</span>
  </div>
  <div class="table-scroll">
    <table class="data-table">
      <thead><tr><th>#</th>{ths}</tr></thead>
      <tbody>{trs}</tbody>
    </table>
  </div>
</div>'''

            fname = _re.sub(r'[^\w\-]', '-', hint.lower())[:30] + '.csv'
            return jsonify({
                "status":       "success",
                "doc_type":     "csv",
                "filename":     fname,
                "mime_type":    "text/csv",
                "content_b64":  None,
                "content_text": csv_text,
                "preview_html": preview_html,
                "size_bytes":   len(csv_text.encode()),
            })

        except Exception as exc:
            import traceback; traceback.print_exc()
            return jsonify({"status": "error", "message": str(exc)}), 500

    # ─────────────────────────────────────────────────────────────────────────
    # XLSX
    # ─────────────────────────────────────────────────────────────────────────
    elif doc_type == 'xlsx':
        try:
            import xlsxwriter, io as _io, base64 as _b64

            xlsx_prompt = f"""Extract all tabular / structured data from the conversation for an Excel spreadsheet.

USER REQUEST: {instruction or 'Create an Excel spreadsheet from the data in this conversation.'}

CONVERSATION:
{conv_text[:6000]}

Return ONLY a JSON object (no markdown fences):
{{
  "filename_hint": "descriptive-name",
  "sheets": [
    {{
      "name": "Sheet Name (max 31 chars)",
      "headers": ["Column1", "Column2", ...],
      "rows": [["val1","val2",...], ...]
    }}
  ]
}}

Create one sheet per logical data group (e.g. separate sheet per year if data spans years).
Include ALL numerical/factual data. Column headers must be clear."""

            result = _llm_provider.chat(
                system="", model=model, temperature=0.15, max_tokens=2500,
                messages=[{"role": "user", "content": xlsx_prompt}],
            )
            raw  = (result.content or "").strip()
            arr_m = _re.search(r'\{[\s\S]*\}', raw)
            struct = json.loads(arr_m.group(0) if arr_m else raw)

            sheets = struct.get('sheets', [])
            hint   = struct.get('filename_hint', topic.lower())

            buf  = _io.BytesIO()
            wb   = xlsxwriter.Workbook(buf, {'in_memory': True})
            hdr_fmt  = wb.add_format({'bold':True,'bg_color':'#1e293b','font_color':'#e2e8f0','border':1,'text_wrap':True})
            cell_fmt = wb.add_format({'border':1,'text_wrap':True})
            alt_fmt  = wb.add_format({'border':1,'bg_color':'#f8fafc','text_wrap':True})

            preview_parts = []
            for sheet in sheets[:5]:
                sname   = str(sheet.get('name', 'Sheet'))[:31]
                headers = sheet.get('headers', [])
                rows_d  = sheet.get('rows', [])
                ws = wb.add_worksheet(sname)
                for j, h in enumerate(headers):
                    ws.write(0, j, str(h), hdr_fmt)
                    ws.set_column(j, j, max(12, len(str(h)) + 4))
                for i, row in enumerate(rows_d):
                    fmt = cell_fmt if i % 2 == 0 else alt_fmt
                    for j, val in enumerate(row[:len(headers)]):
                        ws.write(i+1, j, val, fmt)
                ws.freeze_panes(1, 0)

                # Preview HTML for first sheet
                if not preview_parts:
                    ths = ''.join(f'<th>{_html_esc(h)}</th>' for h in headers)
                    trs = ''.join(
                        '<tr><td class="rn">' + str(i+1) + '</td>' +
                        ''.join(f'<td>{_html_esc(v)}</td>' for v in row) + '</tr>'
                        for i, row in enumerate(rows_d)
                    )
                    preview_parts.append(f'''<div class="table-card">
<div class="table-meta"><span><strong>{len(headers)}</strong> columns</span><span><strong>{len(rows_d)}</strong> rows</span><span>Sheet: {_html_esc(sname)}</span></div>
<div class="table-scroll"><table class="data-table"><thead><tr><th>#</th>{ths}</tr></thead><tbody>{trs}</tbody></table></div></div>''')

            wb.close()
            buf.seek(0)
            b64 = _b64.b64encode(buf.read()).decode()

            fname = _re.sub(r'[^\w\-]', '-', hint.lower())[:30] + '.xlsx'
            return jsonify({
                "status":       "success",
                "doc_type":     "xlsx",
                "filename":     fname,
                "mime_type":    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "content_b64":  b64,
                "content_text": None,
                "preview_html": preview_parts[0] if preview_parts else '',
                "size_bytes":   len(b64)*3//4,
            })

        except Exception as exc:
            import traceback; traceback.print_exc()
            return jsonify({"status": "error", "message": str(exc)}), 500

    # ─────────────────────────────────────────────────────────────────────────
    # MARKDOWN / TXT
    # ─────────────────────────────────────────────────────────────────────────
    elif doc_type in ('markdown', 'txt'):
        try:
            ext = 'md' if doc_type == 'markdown' else 'txt'
            fmt = 'Markdown (with headings, bold, lists, tables)' if doc_type == 'markdown' else 'plain text with clear sections'

            text_prompt = f"""You are a professional writer. Write a {fmt} document based on the conversation.

USER REQUEST: {instruction or f'Write a comprehensive {fmt} document summarising the conversation.'}

CONVERSATION:
{conv_text[:7000]}

Return ONLY the document text — no JSON, no code fences, just the document itself.
Use REAL data, numbers and specific facts from the conversation.
Structure: title → summary → detailed sections → conclusion.
Minimum 400 words."""

            result = _llm_provider.chat(
                system="", model=model, temperature=0.25, max_tokens=2500,
                messages=[{"role": "user", "content": text_prompt}],
            )
            text_content = (result.content or "").strip()

            fname = _re.sub(r'[^\w\-]', '-', topic.lower())[:30] + f'.{ext}'
            mime  = 'text/markdown' if doc_type == 'markdown' else 'text/plain'
            return jsonify({
                "status":       "success",
                "doc_type":     doc_type,
                "filename":     fname,
                "mime_type":    mime,
                "content_b64":  None,
                "content_text": text_content,
                "preview_html": None,
                "size_bytes":   len(text_content.encode()),
            })

        except Exception as exc:
            import traceback; traceback.print_exc()
            return jsonify({"status": "error", "message": str(exc)}), 500

    # ─────────────────────────────────────────────────────────────────────────
    # PDF
    # ─────────────────────────────────────────────────────────────────────────
    elif doc_type == 'pdf':
        try:
            import base64 as _b64
            from services.workspace_export_service import generate_pdf as _gen_pdf

            pdf_prompt = f"""You are a professional report writer. Write a clean, well-structured document in Markdown format based on the conversation below.

USER REQUEST: {instruction or 'Create a professional PDF report from this conversation.'}

CONVERSATION:
{conv_text[:7000]}

Return ONLY the Markdown document text — no code fences, no JSON.
Use real data, numbers, and specific facts from the conversation.
Structure: # Title → ## Summary → ## Key Information → ## Details → ## Conclusion
Use bullet points for lists, bold for key figures, and include all important data."""

            result = _llm_provider.chat(
                system="", model=model, temperature=0.2, max_tokens=2500,
                messages=[{"role": "user", "content": pdf_prompt}],
            )
            md_content = (result.content or "").strip()

            # Extract title from first # heading
            title_match = _re.match(r'^#\s+(.+)', md_content)
            doc_title   = title_match.group(1).strip() if title_match else topic

            _in_tok       = result.usage.prompt_tokens
            _out_tok      = result.usage.completion_tokens
            tokens_used   = result.usage.total_tokens
            try:
                import token_limits as _tl
                if user and user.get('id'):
                    _tl.record_usage(
                        user_id       = user['id'],
                        tokens        = tokens_used,
                        call_type     = 'hub_pdf',
                        agent_name    = agent_name,
                        question      = (instruction or 'pdf generation')[:200],
                        input_tokens  = _in_tok,
                        output_tokens = _out_tok,
                        model         = model,
                    )
            except Exception:
                pass

            pdf_bytes = _gen_pdf(md_content, title=doc_title)
            b64       = _b64.b64encode(pdf_bytes).decode()
            fname     = _re.sub(r'[^\w\-]', '-', topic.lower())[:30] + '.pdf'

            return jsonify({
                "status":       "success",
                "doc_type":     "pdf",
                "filename":     fname,
                "mime_type":    "application/pdf",
                "content_b64":  b64,
                "content_text": None,
                "preview_html": None,
                "size_bytes":   len(pdf_bytes),
            })

        except Exception as exc:
            import traceback; traceback.print_exc()
            return jsonify({"status": "error", "message": str(exc)}), 500

    else:
        return jsonify({"status": "error", "message": f"Unsupported type: {doc_type}"}), 400


def _html_esc(s):
    """Escape &, <, > for safe inclusion in generated HTML preview fragments."""
    return str(s or '').replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
